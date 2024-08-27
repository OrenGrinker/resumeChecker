from PyPDF2 import PdfReader
from docx import Document
from openai import OpenAI

def process_resumes(uploaded_files):
    candidates = []
    for uploaded_file in uploaded_files:
        if uploaded_file.type == "application/pdf":
            text = extract_text_from_pdf(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = extract_text_from_docx(uploaded_file)
        else:
            continue
        candidates.append({"name": uploaded_file.name, "text": text})
    return candidates

def extract_text_from_pdf(uploaded_file):
    reader = PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text


def extract_text_from_docx(uploaded_file):
    doc = Document().Document(uploaded_file)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text


def find_suitable_candidates(client, candidates, query):
    suitable_candidates = []

    for candidate in candidates:
        # Construct the content with text
        user_content = [
            {
                "type": "text",
                "text": f"Given the following resume, determine if this candidate is suitable for the following query: {query}. "
                        f"Please respond with 'Suitable' if the candidate matches the requirements or 'Not suitable' if they do not. "
                        f"If suitable, also provide a brief summary of relevant experience and include the candidate's contact information."
            },
            {
                "type": "text",
                "text": candidate["text"]
            }
        ]

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "user",
                    "content": user_content
                }
            ],
            temperature=1,
            max_tokens=3000,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0,
            response_format={
                "type": "text"
            }
        )

        summary = response.choices[0].message.content.strip()

        if "Suitable" in summary:
            # Extract contact information from the summary
            contact_info = extract_contact_info(summary)

            suitable_candidates.append({
                "name": candidate["name"],
                "summary": summary,
                "contact_info": contact_info,
                "text": candidate["text"]
            })

    return suitable_candidates


def extract_contact_info(summary):
    # This is a simple placeholder for whatever logic you might use to extract the contact info.
    # You would need to parse the summary, look for patterns like email addresses, phone numbers, etc.
    # For now, we'll just return the summary as a placeholder.
    contact_info = "Extracted contact info"
    return contact_info


